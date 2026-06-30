from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "คู่มือโครงสร้างและจุดแก้ไขระบบจองห้อง_ฉบับสรุป.docx"
FONT = "Leelawadee UI"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(90, 98, 108)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    run.font.name = FONT
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "FFF4CE")
    set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        header.cells[i].width = Inches(widths[i])
        header.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(header.cells[i], LIGHT_BLUE)
        set_cell_margins(header.cells[i])
        p = header.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.text = "คู่มือระบบจองห้อง | Room Reservation"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.runs[0].font.name = FONT
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = MUTED
add_page_number(section.footer.paragraphs[0])

# Editorial cover
doc.add_paragraph().paragraph_format.space_after = Pt(80)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(16)
r = kicker.add_run("คู่มืออ้างอิงฉบับย่อ")
r.bold = True
r.font.name = FONT
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(BLUE)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
r = title.add_run("โครงสร้างและจุดแก้ไข\nระบบจองห้อง")
r.bold = True
r.font.name = FONT
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(90)
r = subtitle.add_run("อธิบายหน้าที่ของแต่ละหน้า ไฟล์ที่เกี่ยวข้อง และบรรทัดสำคัญ")
r.font.name = FONT
r.font.size = Pt(13)
r.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("โปรเจกต์: room-reservation\n").bold = True
meta.add_run("Frontend: React  |  Backend: Express  |  Database: MySQL\n")
meta.add_run("อ้างอิงจากโค้ด ณ วันที่ 25 มิถุนายน 2026")

doc.add_page_break()

add_heading(doc, "1. ภาพรวมระบบ", 1)
doc.add_paragraph(
    "ระบบแบ่งออกเป็น Frontend สำหรับหน้าเว็บไซต์, Backend สำหรับ API และ MySQL "
    "สำหรับจัดเก็บข้อมูล โดยลำดับการทำงานคือ ผู้ใช้ → React → API → MySQL"
)
add_table(
    doc,
    ["ส่วน", "ตำแหน่ง", "หน้าที่"],
    [
        ("Frontend", "frontend/src", "หน้าเว็บไซต์ React, การแสดงผล และการเรียก API"),
        ("Backend", "backend", "Express API, ตรวจสอบข้อมูล และคำสั่งฐานข้อมูล"),
        ("Database", "backend/db.js", "ตั้งค่าการเชื่อมต่อ MySQL"),
        ("เส้นทางหน้า", "frontend/src/App.js:23-125", "กำหนด URL และ Component ของทุกหน้า"),
    ],
    [1.15, 2.05, 3.3],
)

add_note(
    doc,
    "จุดที่ทำให้สับสน",
    "ชื่อไฟล์ Home สลับความหมายกัน: adminHome.js คือหน้าผู้ใช้ /home "
    "ส่วน userHome.js คือหน้าแอดมิน /admin การจับคู่อยู่ที่ App.js บรรทัด 9-10",
)

add_heading(doc, "2. หน้าเว็บไซต์ทั้งหมด", 1)
pages = [
    ("/", "หน้าแรกสาธารณะ ค้นหาและดูปฏิทิน", "pages/PublicHome.js:51"),
    ("/login", "เข้าสู่ระบบ", "pages/Login.js:5"),
    ("/register", "สมัครสมาชิก @lru.ac.th", "Register.js:6"),
    ("/verify-email", "ยืนยันอีเมลด้วย token", "VerifyEmail.js:6"),
    ("/forgot-password", "ขอลิงก์เปลี่ยนรหัสผ่าน", "ForgotPassword.js:6"),
    ("/reset-password", "ตั้งรหัสผ่านใหม่", "ResetPassword.js:6"),
    ("/home", "หน้าผู้ใช้ ค้นหาและจองห้อง", "pages/adminHome.js:169"),
    ("/reservation-list", "รายการจองของผู้ใช้", "pages/ReservationList.js:66"),
    ("/profile", "แก้ไขข้อมูลส่วนตัว", "pages/Profile.js:6"),
    ("/report", "รายงานฝั่งผู้ใช้", "pages/Report.js:38"),
    ("/admin", "Dashboard แอดมิน", "pages/userHome.js:39"),
    ("/admin/reservations", "อนุมัติและปฏิเสธการจอง", "pages/ReservationList.js:66"),
    ("/admin/users", "จัดการผู้ใช้และสิทธิ์", "pages/UserManage.js:6"),
    ("/admin/rooms", "จัดการอาคาร ประเภทห้อง และห้อง", "pages/BuildingRoomAdmin.js:8"),
    ("/admin/schedules", "จัดการตารางเรียน", "pages/ScheduleManage.js:213"),
    ("/admin/report", "รายงานฝั่งแอดมิน", "pages/Report.js:38"),
]
add_table(doc, ["URL", "หน้าที่", "ไฟล์หลัก : บรรทัด"], pages, [1.45, 2.8, 2.25])

add_heading(doc, "3. จุดแก้ไข Frontend ที่สำคัญ", 1)
add_heading(doc, "3.1 เมนูและสิทธิ์เข้าใช้งาน", 2)
add_bullet(doc, "เพิ่ม/ลบ/เปลี่ยนชื่อเมนู: components/MainMenuTabs.js:14-50")
add_bullet(doc, "ปรับสีและรูปแบบเมนู: components/MainMenuTabs.css")
add_bullet(doc, "กำหนดหน้าสำหรับผู้ใช้: routes/PrivateRoute.js:3-16")
add_bullet(doc, "กำหนดหน้าสำหรับแอดมิน: routes/AdminRoute.js:3-18")
add_bullet(doc, "เพิ่มหรือลบ URL ของหน้า: App.js:25-125")

add_heading(doc, "3.2 หน้าจองห้องของผู้ใช้", 2)
add_table(
    doc,
    ["สิ่งที่ต้องการแก้", "ตำแหน่ง"],
    [
        ("โหลดห้อง อาคาร ประเภทห้อง และรายการจอง", "pages/adminHome.js:234-238"),
        ("เงื่อนไขค้นหาห้องว่าง", "pages/adminHome.js:324-410"),
        ("การกดวันที่ในปฏิทิน", "pages/adminHome.js:461"),
        ("ฟอร์มเลือกห้องและเวลา", "pages/adminHome.js:486-575"),
        ("ส่งคำขอจองเข้า API", "pages/adminHome.js:576-640"),
        ("หน้าตาและข้อความของหน้า", "pages/adminHome.js:641-948"),
        ("สี ระยะห่าง และ Responsive", "pages/Home.css"),
    ],
    [3.4, 3.1],
)

add_heading(doc, "3.3 ปฏิทิน", 2)
add_bullet(doc, "เงื่อนไขว่ารายการใดต้องแสดง: components/Calendar.js:12-219")
add_bullet(doc, "โครงสร้างปฏิทิน: components/Calendar.js:220-234")
add_bullet(doc, "สีของตารางเรียน การจอง และวันที่: components/Calendar.css")
doc.add_paragraph("Calendar Component นี้ถูกใช้ร่วมกันในหน้าสาธารณะ หน้าผู้ใช้ และหน้าแอดมิน")

add_heading(doc, "3.4 รายการจองและการอนุมัติ", 2)
add_table(
    doc,
    ["การทำงาน", "ตำแหน่ง"],
    [
        ("โหลดรายการจอง", "pages/ReservationList.js:76-112"),
        ("ผู้ใช้ยกเลิกรายการ", "pages/ReservationList.js:122-153"),
        ("แอดมินเปลี่ยนสถานะ", "pages/ReservationList.js:155-193"),
        ("สร้างไฟล์ PDF", "pages/ReservationList.js:234-355"),
        ("พิมพ์เอกสาร", "pages/ReservationList.js:366-600"),
        ("ปุ่มอนุมัติ ปฏิเสธ ยกเลิก และพิมพ์", "pages/ReservationList.js:713-786"),
        ("รูปแบบหน้ารายการ", "pages/ReservationList.css"),
    ],
    [3.1, 3.4],
)

add_heading(doc, "3.5 อาคาร ประเภทห้อง และห้อง", 2)
add_bullet(doc, "หน้ารวมและปุ่มสลับแท็บ: pages/BuildingRoomAdmin.js:20-47")
add_bullet(doc, "เพิ่ม/แก้ไข/เปลี่ยนสถานะอาคาร: pages/admin/BuildingsTab.js:8-135")
add_bullet(doc, "เพิ่ม/แก้ไขประเภทห้อง: pages/admin/RoomTypesTab.js:4-128")
add_bullet(doc, "ห้อง ชั้น ความจุ อุปกรณ์ และสถานะ: pages/admin/RoomsTab.js:13-260")
add_bullet(doc, "รูปแบบหน้าทั้งหมด: pages/admin/AdminRooms.css")

add_heading(doc, "3.6 ตารางเรียน ผู้ใช้ รายงาน และโปรไฟล์", 2)
add_table(
    doc,
    ["หน้า", "จุดสำคัญ"],
    [
        ("ตารางเรียน", "ScheduleManage.js:123-212 การทำซ้ำ; 316 ลบ; 323-373 เพิ่ม/แก้ไข"),
        ("ผู้ใช้", "UserManage.js:31 โหลด; 47 เปลี่ยนสิทธิ์; 85 เพิ่ม; 134 แก้รหัสผ่าน"),
        ("รายงาน", "Report.js:54-56 โหลดข้อมูล; 125-309 สรุปและตารางรายงาน"),
        ("โปรไฟล์", "Profile.js:31 โหลดข้อมูล; 60-108 บันทึกการแก้ไข"),
    ],
    [1.4, 5.1],
)

add_heading(doc, "4. Backend และ API", 1)
doc.add_paragraph(
    "Backend เริ่มทำงานจาก backend/server.js และรวมทุก Route ไว้ใต้ /api "
    "โดยใช้พอร์ต 3002 ที่ server.js บรรทัด 59"
)
add_table(
    doc,
    ["งาน", "ไฟล์ Backend", "จุดเริ่ม"],
    [
        ("สมัคร/เข้าสู่ระบบ/ยืนยันอีเมล", "routes/login.js", "158, 262, 304, 357, 403"),
        ("รายการจองและสถานะ", "routes/reservations.js", "5, 63, 193, 227"),
        ("ห้องและค้นหาห้องว่าง", "routes/rooms.js", "47, 57, 143, 227"),
        ("อาคาร", "routes/buildings.js", "8, 21, 48"),
        ("ประเภทห้อง", "routes/roomTypes.js", "14, 24, 41"),
        ("ตารางเรียน", "routes/schedules.js", "38, 95, 106, 161"),
        ("ผู้ใช้และสิทธิ์", "routes/users.js", "7, 32, 96, 127, 186, 213"),
    ],
    [2.35, 2.4, 1.75],
)

add_heading(doc, "4.1 จุดตรวจเวลาห้องชน", 2)
add_bullet(doc, "ตรวจชนกับตารางเรียน: backend/routes/reservations.js:89-115")
add_bullet(doc, "ตรวจชนกับรายการจองอื่น: backend/routes/reservations.js:117-143")
add_bullet(doc, "ค้นหาห้องว่างจากทั้งตารางเรียนและรายการจอง: backend/routes/rooms.js:57-139")
add_bullet(doc, "สถานะที่แอดมินอนุญาตให้เปลี่ยน: backend/routes/reservations.js:230")

add_heading(doc, "4.2 ฐานข้อมูล", 2)
add_table(
    doc,
    ["สิ่งที่แก้", "ตำแหน่ง"],
    [
        ("Host, user, password และชื่อฐานข้อมูล", "backend/db.js:3-8"),
        ("พอร์ต Backend", "backend/server.js:59"),
        ("URL Frontend ในอีเมล", "backend/routes/login.js:7"),
        ("SMTP สำหรับส่งอีเมล", "ไฟล์ .env และ backend/routes/login.js:75-92"),
    ],
    [3.0, 3.5],
)

add_heading(doc, "5. ตารางฐานข้อมูลหลัก", 1)
add_table(
    doc,
    ["ตาราง", "เก็บข้อมูล"],
    [
        ("users", "ผู้ใช้ รหัสผ่าน สิทธิ์ เบอร์โทร และสถานะยืนยันอีเมล"),
        ("buildings", "ชื่ออาคารและสถานะ"),
        ("room_types", "ประเภทห้องและสถานะ"),
        ("rooms", "ห้อง ชั้น ความจุ อาคาร ประเภท อุปกรณ์ และสถานะ"),
        ("schedules", "ตารางเรียน ห้อง วันที่ เวลา และรูปแบบการทำซ้ำ"),
        ("reservations", "ผู้จอง ห้อง วันเวลา จำนวนคน วัตถุประสงค์ และสถานะ"),
    ],
    [1.55, 4.95],
)

add_heading(doc, "6. สิ่งที่ควรแก้ก่อนนำขึ้นใช้งานจริง", 1)
security_rows = [
    ("เร่งด่วน", "รหัสผ่านเก็บเป็นข้อความธรรมดา", "ใช้ bcrypt hash และ compare แทนการค้นหา password ตรง ๆ"),
    ("เร่งด่วน", "Backend ไม่มีระบบ token ตรวจสิทธิ์", "เพิ่ม JWT/session middleware ให้ API สำคัญ"),
    ("สูง", "รหัสผ่าน MySQL อยู่ใน db.js", "ย้ายค่าทั้งหมดไป .env"),
    ("สูง", "API URL กระจายหลายไฟล์", "สร้างไฟล์ config หรือใช้ REACT_APP_API_URL"),
    ("กลาง", "ชื่อ adminHome/userHome สลับความหมาย", "เปลี่ยนชื่อเป็น UserHome และ AdminHome"),
    ("กลาง", "มีไฟล์เก่าที่ดูเหมือนไม่ได้ใช้งาน", "ตรวจ import ก่อนลบ AdminSchedule, ReservationTable และเมนูเก่า"),
]
add_table(doc, ["ระดับ", "ปัญหา", "แนวทางแก้"], security_rows, [0.85, 2.35, 3.3])

add_note(
    doc,
    "หมายเหตุเรื่องเลขบรรทัด",
    "เลขบรรทัดทั้งหมดอ้างอิงจากโค้ดปัจจุบัน หากเพิ่มหรือลบโค้ด เลขบรรทัดอาจเลื่อน "
    "ควรค้นหาด้วยชื่อฟังก์ชันหรือข้อความ API ควบคู่กัน",
)

add_heading(doc, "7. สรุปจุดที่แก้บ่อยที่สุด", 1)
add_table(
    doc,
    ["ต้องการแก้", "ไปที่ไฟล์"],
    [
        ("URL หรือเพิ่มหน้าใหม่", "frontend/src/App.js"),
        ("เมนูด้านบน", "frontend/src/components/MainMenuTabs.js"),
        ("หน้าจองห้อง", "frontend/src/pages/adminHome.js"),
        ("ปฏิทิน", "frontend/src/components/Calendar.js และ Calendar.css"),
        ("รายการจอง/อนุมัติ", "frontend/src/pages/ReservationList.js"),
        ("ข้อมูลห้องและอาคาร", "frontend/src/pages/admin/*.js และ backend/routes/rooms.js"),
        ("ตารางเรียน", "frontend/src/pages/ScheduleManage.js และ backend/routes/schedules.js"),
        ("Login/Register/Email", "backend/routes/login.js"),
        ("ฐานข้อมูล", "backend/db.js"),
    ],
    [2.35, 4.15],
)

doc.core_properties.title = "คู่มือโครงสร้างและจุดแก้ไขระบบจองห้อง"
doc.core_properties.subject = "สรุปหน้าเว็บไซต์ ไฟล์ และบรรทัดสำคัญ"
doc.core_properties.author = "Codex"
doc.save(OUTPUT)
print(OUTPUT)
